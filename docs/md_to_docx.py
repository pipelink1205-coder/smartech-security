#!/usr/bin/env python3
"""Convierte borradores-servicios.md a Word (.docx)."""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Instala python-docx: pip install python-docx")
    sys.exit(1)


def strip_md_bold(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def add_rich_paragraph(doc: Document, text: str, style: str | None = None):
    """Párrafo con negritas **texto**."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)
    return p


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_blockquote = False

    for raw in lines:
        line = raw.rstrip()

        if line.strip() == "---":
            doc.add_page_break()
            continue

        if not line.strip():
            in_blockquote = False
            continue

        # Tabla separador
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue

        # Tabla fila
        if line.strip().startswith("|") and line.count("|") >= 2:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            add_rich_paragraph(doc, "  •  ".join(strip_md_bold(c) for c in cells if c))
            continue

        # Blockquote
        if line.startswith("> "):
            add_rich_paragraph(doc, line[2:], style="Intense Quote")
            in_blockquote = True
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(strip_md_bold(line[2:]), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md_bold(line[3:]), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(strip_md_bold(line[4:]), level=2)
            continue

        # Lista numerada
        m_num = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m_num:
            add_rich_paragraph(doc, f"{m_num.group(1)}. {m_num.group(2)}")
            continue

        # Lista viñeta
        if line.startswith("- "):
            add_rich_paragraph(doc, line[2:], style="List Bullet")
            continue

        # Checkbox checklist
        if line.startswith("- [ ]"):
            add_rich_paragraph(doc, "☐ " + line[6:], style="List Bullet")
            continue

        add_rich_paragraph(doc, line)

    doc.save(docx_path)
    print(f"Generado: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).resolve().parent
    md = base / "borradores-servicios.md"
    out = base / "borradores-servicios.docx"
    if not md.exists():
        print(f"No existe: {md}")
        sys.exit(1)
    convert(md, out)
