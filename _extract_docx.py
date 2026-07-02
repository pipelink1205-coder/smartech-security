from docx import Document
from pathlib import Path

doc = Document(Path(r'c:\Users\Diana\Downloads\borradores-servicios_listo.docx'))
for para in doc.paragraphs:
    t = para.text.strip()
    if t:
        print(t)
print('---TABLES---')
for ti, table in enumerate(doc.tables):
    print(f'[TABLE {ti}]')
    for row in table.rows:
        print(' | '.join(cell.text.strip().replace('\n', ' ') for cell in row.cells))
